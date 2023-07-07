from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

class Textsdoc:
    def __init__(self):
        pass
    
    def add_heading(self, document, text, level):
        heading = document.add_heading(text, level)
        heading.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def add_paragraph(self, document, text):
        paragraph = document.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        font = paragraph.style.font
        font.name = 'Arial'
        font.size = Pt(12)
    
    def text2(self, document):
        self.add_heading(document, '1. Introdução', 1)
        self.add_paragraph(document, 'A linearidade de um procedimento analítico é a sua capacidade de obter resultados que sejam diretamente proporcionais à concentração de um analito em uma amostra.')
    
    def text3(self, document):
        self.add_heading(document, '2. Coleta de Dados', 1)
        self.add_paragraph(document, 'A seguir, apresentam-se os dados coletados:')
    
    def text4(self, document):
        self.add_heading(document, '3. Método dos Mínimos Quadrados Ordinários Estimação', 1)
        self.add_paragraph(document, 'O método dos mínimos quadrados é uma eficiente estratégia de estimação dos parâmetros da regressão e sua aplicação não é limitada apenas às relações lineares. Nesta seção utilizou-se o Método dos Mínimos Quadrados Ordinários.')
        
        self.add_heading(document, '3.1. Teste do coeficiente angular', 2)
        self.add_paragraph(document, 'Para avaliar a significância do modelo utilizou-se o teste F da ANOVA. Neste caso, testou-se as hipóteses:')
        self.add_paragraph(document, 'H0: coeficiente angular igual a zero;')
        self.add_paragraph(document, 'H1: coeficiente angular diferente de zero.')
    
    def text5(self, document, p_value):
        if p_value > 0.05:
            self.add_paragraph(document, f'Como P-valor ({p_value}) do teste ANOVA é maior que 0,05 (conforme especificado), não rejeita-se a hipótese nula (intercepto igual ao zero) ao nível de significância de 5%. Logo, conclui-se que o intercepto é estatísticamente igual ao zero.')
        else:
            self.add_paragraph(document, f'Como P-valor ({p_value}) do teste ANOVA é menor ou igual a 0,05 (conforme especificado), rejeita-se a hipótese nula (intercepto igual ao zero) ao nível de significância de 5%. Logo, conclui-se que o intercepto é estatísticamente diferente de zero.')
