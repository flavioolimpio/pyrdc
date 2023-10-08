import streamlit as st
from streamlit_option_menu import option_menu
import pandas as pd
import numpy as np
import statsmodels.api as sm
import statsmodels.formula.api as smf
from texts import Texts
from testes import testes_variancia
from plots import Plots
import os
from io import BytesIO
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from texts2doc import Textsdoc
#from normalidade import test_norm

def home():
    st.title("Página Inicial")
    gettext = Texts()
    
    text1 = gettext.text1()
    st.markdown('#pyRDC')
    st.markdown('{}'.format(text1), unsafe_allow_html=True)


def add_df_to_word(document, df):
    # Adicione uma tabela com o número de linhas e colunas do DataFrame
    table = document.add_table(rows=df.shape[0]+1, cols=df.shape[1])
    
    # Adicione os cabeçalhos da tabela
    for i, column in enumerate(df.columns):
        table.cell(0, i).text = column
    
    # Adicione os dados da tabela
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            table.cell(i+1, j).text = str(df.iloc[i, j])


def add_df_to_word(document, df, caption):
    # Adicione a legenda da tabela
    texts = Textsdoc()
    texts.add_paragraph(document, caption)
    
    # Adicione uma tabela com o número de linhas e colunas do DataFrame
    table = document.add_table(rows=df.shape[0]+1, cols=df.shape[1])
    
    # Adicione os cabeçalhos da tabela
    for i, column in enumerate(df.columns):
        cell = table.cell(0, i)
        cell.text = column
        
        # Adicione bordas superior e inferior ao cabeçalho da tabela
        tcPr = cell._element.tcPr
        tcBorders = parse_xml(r'<w:tcBorders %s><w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/></w:tcBorders>' % nsdecls('w'))
        tcPr.append(tcBorders)
    
    # Adicione os dados da tabela
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            cell = table.cell(i+1, j)
            
            value = df.iloc[i, j]
            
            # Arredonde os valores numéricos para 4 casas decimais
            if isinstance(value, float):
                value = round(value, 4)
            
            cell.text = str(value)
            
            # Adicione uma borda inferior à última linha da tabela
            if i == df.shape[0]-1:
                tcPr = cell._element.tcPr
                tcBorders = parse_xml(r'<w:tcBorders %s><w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/></w:tcBorders>' % nsdecls('w'))
                tcPr.append(tcBorders)

                tcBorders = parse_xml(r'<w:tcBorders %s><w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/></w:tcBorders>' % nsdecls('w'))

def save_doc():
    
    document = Document()
    
    texts = Textsdoc()
    
    data = pd.read_csv('Analise_teste_hetero.csv', sep=',')
    
    testes = testes_variancia()
    
    aov_table, coef_table, conf_int, pearson_table = testes.anova(data)
    
    bp_test, gq_test, bg_test, teste_shapiro = testes.homoscedasticity_test(data)
    
    texts.text2(document)
    texts.text3(document)
    add_df_to_word(document, data, 'Tabela 1: Conjunto de dados para o estudo de Linearidade.')
    texts.text4(document)
    add_df_to_word(document, aov_table, "Tabela 2: Esta tabela apresenta os resultados da análise de variância (ANOVA) para avaliar a significância das diferenças entre as médias dos grupos." +  
    "A tabela inclui a fonte de variação, os graus de liberdade, a soma dos quadrados, a média dos quadrados, o valor F e o valor P.")
    add_df_to_word(document, coef_table, "Tabela 3: Parâmetros dos coeficientes.")
    add_df_to_word(document, conf_int, "Tabela 4: Intervalo de confiança para os parâmetros.")
    
    p_value = 0.03
    texts.text5(document, p_value)
    add_df_to_word(document, pearson_table, "Tabela 5: Esta tabela apresenta o resultado do cálculo do coeficiente de correlação de Pearson entre duas variáveis." + 
    "O coeficiente mede a força e a direção da relação linear entre as variáveis.")
    
    #add_df_to_word(document, bp_test)
    #add_df_to_word(document, gq_test)
    #add_df_to_word(document, bg_test)
    #add_df_to_word(document, teste_shapiro)
    
    #document.save('Relatorio_v03.docx')
    
    return document


def linearidade():
    st.title("Resultado da Análise")
    file = st.file_uploader("Arraste e solte um arquivo .csv ou .xlsx aqui", type=['csv', 'xlsx'])
    if file:
        if file.type == 'text/csv':
            first_line = file.readline().decode('utf-8')
            sep = ',' if ',' in first_line else ';'
            file.seek(0)
            df = pd.read_csv(file, sep=sep)
        else:
            df = pd.read_excel(file)

        gettext = Texts()

        text2 = gettext.text2()
        st.markdown('{}'.format(text2), unsafe_allow_html=True)

        text3 = gettext.text3()
        st.markdown('{}'.format(text3), unsafe_allow_html=True)
        st.write("Tabela 1: Conjunto de dados para o estudo de Linearidade.")
        st.table(df)

        text4 = gettext.text4()
        st.markdown('{}'.format(text4), unsafe_allow_html=True)

        gettestes = testes_variancia()

        aov_table, coef_table, conf_int, pearson_table = gettestes.anova(data=df)

        st.write("Tabela 2: Esta tabela apresenta os resultados da análise de variância (ANOVA) para avaliar a significância das diferenças entre as médias dos grupos." +  
                 "A tabela inclui a fonte de variação, os graus de liberdade, a soma dos quadrados, a média dos quadrados, o valor F e o valor P.")
        st.table(aov_table)

        p_value_anova = aov_table['P-valor'][0]

        text5 = gettext.text5(p_value=p_value_anova)
        st.markdown('{}'.format(text5), unsafe_allow_html=True)

        st.write("As estimativas dos parâmetros (coeficientes de regressão) e seus respectivos intervalos de confiança de 95% são:")
        st.write("Tabela 3: Parâmetros dos coeficientes.")
        st.table(coef_table)

        st.write("Tabela 4: Intervalo de confiança para os parâmetros.")
        st.table(conf_int)

        p_value_coef = coef_table['P-valor'][0]
        text6 = gettext.text6(p_value=p_value_coef)
        st.markdown('{}'.format(text6), unsafe_allow_html=True)

        cc = pearson_table['Coeficiente de Correlação']
        r_value = cc.item()
        text7 = gettext.text7(r_value=r_value)
        st.markdown('{}'.format(text7), unsafe_allow_html=True)
        st.write("Tabela 5: Esta tabela apresenta o resultado do cálculo do coeficiente de correlação de Pearson entre duas variáveis." + 
        "O coeficiente mede a força e a direção da relação linear entre as variáveis.")
        st.table(pearson_table)

        getfigs = Plots()
        x = df.columns[0]
        y = df.columns[1]
        z = df.columns[3]
        fig_residuo, fig_residuo_pad, fig_qqplot, fig_coleta = getfigs.graficos(df[x].values,df[y].values, df[z].values)

        text8 = gettext.text8()
        st.markdown('{}'.format(text8), unsafe_allow_html=True)

        st.write("Figura 1: Gráfico de Resíduos Padronizados vs Valores Ajustados.")
        st.pyplot(fig=fig_residuo)
        st.write("Figura 2: Gráfico de Probabilidade Normal dos Resíduos.")
        st.pyplot(fig=fig_residuo_pad)
        st.write("Figura 3: Gráfico de Resíduos vs Valores Ajustados.")
        st.pyplot(fig=fig_qqplot)
        st.write("Figura 4: Gráfico de Resíduos vs Ordem de Coleta")
        st.pyplot(fig=fig_coleta)

        bp_test, gq_test, bg_test, shapiro = gettestes.homoscedasticity_test(data=df)

        ps = shapiro['P-valor']
        p_value_shapiro = ps.item()
        text9 = gettext.text9(p_value=p_value_shapiro)
        st.markdown('{}'.format(text9), unsafe_allow_html=True)
        st.write("Tabela 6: Teste de normalidade de Shapiro-Wilk com 95%.")
        st.table(shapiro)

        pb = bp_test['P-valor']
        p_value_bp = pb.item()
        text10 = gettext.text10(p_value=p_value_bp)
        st.markdown('{}'.format(text10), unsafe_allow_html=True)
        st.write("Tabela 7: Teste de homecedasticidade Breusch-Pagan")
        st.table(bp_test)

        pbg = bg_test['P-valor']
        p_value_bpg = pbg.item()
        text11 = gettext.text11(p_value=p_value_bpg)
        st.markdown('{}'.format(text11), unsafe_allow_html=True)
        st.write("Tabela 8: Teste de independência de Breusch-Godfrey")
        st.table(bg_test)
        # Create a button that will open a dialog box to choose the path for the doc file.
        doc = save_doc()
        with BytesIO() as output:
            doc.save(output)
            data = output.getvalue()

        st.download_button(
            label='Baixar relatório',
            data=data,
            file_name='report.docx',
            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        
        #st.write("Teste de homecedasticidade Goldfeld-Quandt")
        #st.table(gq_test)

def precisao():
    pass

def exatidao():
    pass

def robustez():
    pass
     

# Criação das páginas
PAGES = {
    "Home": home,
    "Linearidade": linearidade,
    "Precisão": precisao, 
    "Exatidão": exatidao, 
    'Robustez': robustez
}

# Criação do menu suspenso com ícones
with st.sidebar:
    selected = option_menu("Main Menu", ["Home", 'Linearidade', 'Precisão', 'Exatidão', 'Robustez'], icons=['house', 'graph-up', 'target', 'check-circle', 'shield'], menu_icon="cast", default_index=1)
    selected

# Renderização da página selecionada
page = PAGES[selected]
page()
